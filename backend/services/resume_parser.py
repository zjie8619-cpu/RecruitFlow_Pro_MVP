# backend/services/resume_parser.py
import os
import re
import textwrap
from pathlib import Path
from typing import Dict, List, Tuple

import chardet
import fitz
import pandas as pd
import pytesseract
from PIL import Image
from pdf2image import convert_from_path

# optional imports (best-effort)
try:  # pragma: no cover
    import pdfplumber  # type: ignore
except Exception:  # pragma: no cover
    pdfplumber = None

try:  # pragma: no cover
    import docx2txt  # type: ignore
except Exception:  # pragma: no cover
    docx2txt = None

try:  # pragma: no cover
    from docx import Document  # type: ignore
except Exception:  # pragma: no cover
    Document = None

try:  # pragma: no cover
    from backend.services.ai_client import get_client_and_cfg, chat_completion
except Exception:  # pragma: no cover
    get_client_and_cfg = None
    chat_completion = None

import zipfile
from xml.etree import ElementTree

default_tesseract = Path(r"C:\Program Files\Tesseract-OCR\tesseract.exe")
if default_tesseract.exists():
    pytesseract.pytesseract.tesseract_cmd = str(default_tesseract)

SUPPORTED_EXT = {".pdf", ".docx", ".txt", ".jpg", ".jpeg", ".png"}


def _clean_text(text: str) -> str:
    text = text.replace("\x00", " ").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _detect_encoding_and_read(raw: bytes) -> str:
    enc = chardet.detect(raw).get("encoding") or "utf-8"
    try:
        return raw.decode(enc, errors="ignore")
    except Exception:
        return raw.decode("utf-8", errors="ignore")


def extract_contacts(text: str) -> Dict[str, str]:
    """提取联系方式，兼容含空格/连字符的手机号与被分隔的邮箱"""
    clean_text = text or ""
    # 邮箱：允许中间有空格或换行，通过移除空白后再匹配
    email = ""
    compact_email_text = re.sub(r"\s+", "", clean_text)
    email_match = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", compact_email_text)
    if email_match:
        email = email_match.group(0)
    else:
        email_match = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", clean_text)
        if email_match:
            email = email_match.group(0)

    # 手机号：允许空格、短横线、点号作为分隔
    phone = ""
    phone_pattern = re.compile(r"(?:\+?86[-\s.]*)?(1[3-9]\d)[-\s.]*([\d][-\s.]?){8}")
    match = phone_pattern.search(clean_text)
    if match:
        digits = re.sub(r"[^\d]", "", match.group(0))
        if digits.startswith("86") and len(digits) > 11:
            digits = digits[-11:]
        phone = digits
    else:
        phone_match = re.search(r"(?<!\d)(1[3-9]\d{9})(?!\d)", clean_text)
        if phone_match:
            phone = phone_match.group(1)

    return {"email": email.strip(), "phone": phone.strip()}


NAME_PATTERNS = [
    re.compile(r"(?:姓\s*名|姓名|Name|name)[:：\s]*([\u4e00-\u9fa5·]{2,8})"),
    re.compile(r"^([\u4e00-\u9fa5·]{2,8})[，,。\s]*(?:男|女)\b", re.MULTILINE),
    re.compile(r"^\s*Name[:：\s]*([A-Za-z][A-Za-z\s\.\-]{1,30})", re.MULTILINE),
    re.compile(r"(?:我叫|我是|本人姓)([\u4e00-\u9fa5·]{2,4})"),
    re.compile(r"(?:My name is|I am)\s+([A-Za-z][A-Za-z\s\.\-]{1,30})", re.IGNORECASE),
]

NAME_STOP_WORDS = {
    # 常见职位/通用词
    "课程顾问",
    "顾问",
    "老师",
    "班主任",
    "岗位",
    "职位",
    "简历",
    "应聘",
    "求职",
    "个人简历",
    "面试",
    "面销",
    "机构",
    "尚德机构",
    "成长经历",
    "联系方法",
    "联系方式",
    "联系方式：",
    "联系方式:",
    "联系人",
    "联系人信息",
    # 基础信息字段
    "性别",
    "男",
    "女",
    "年龄",
    "年龄段",
    "生日",
    "出生年月",
    "出生日期",
    "出生年",
    "身高",
    "体重",
    "民族",
    "籍贯",
    "户籍",
    "现居地",
    "现住址",
    "家庭住址",
    "婚姻状况",
    "政治面貌",
    "求职意向",
    "职业目标",
    # 联系方式字段
    "电话",
    "手机",
    "手机号",
    "座机",
    "微信",
    "微信号",
    "微信号：",
    "微信号:",
    "微信：",
    "微信:",
    "QQ",
    "邮箱",
    "电子邮箱",
    # 教育/经历
    "教育经历",
    "教育背景",
    "学历",
    "最高学历",
    "大专",
    "本科",
    "硕士",
    "博士",
    "实习经历",
    "工作经历",
    "项目经历",
    "实习经验",
    "年经验",
    "专业资质",
    # 其他噪声字段
    "掌握技能",
    "技能特长",
    "资格证书",
    "证书",
    "荣誉奖项",
    "兴趣爱好",
    "自我评价",
    "亮点",
    "摘要",
    "简介",
    "简述",
    "特点",
    "职业标签",
    "核心竞争力",
    "优势总结",
    "责任描述",
    "掌握语言",
    "期望薪资",
    "紧急联系人",
    "备注",
}

NAME_DISALLOWED_SUBSTRINGS = {
    "课程",
    "规划",
    "机构",
    "方式",
    "渠道",
    "顾问",
    "老师",
    "教师",
    "教练",
    "班主任",
    "校长",
    "院长",
    "主任",
    "经理",
    "主管",
    "教育",
    "经历",
    "背景",
    "信息",
    "简历",
    "职位",
    "岗位",
    "招聘",
    "面试",
    "面销",
    "求职",
    "联系",
    "项目",
    "完成",
    "成交",
    "全国",
    "学生",
    "学士",
    "硕士",
    "博士",
    "中学",
    "小学",
    "教师资格",
    "毕业",
    "大学",
    "学院",
    "经验",
    "技能",
    "证书",
    "语言",
    "薪资",
    "住址",
    "现居",
    "邮箱",
    "电话",
    "微信",
    "QQ",
    "链接",
    "主页",
    "网站",
    "公众号",
    "教育部",
}

SUSPICIOUS_NAME_KEYWORDS = {
    "奖",
    "获奖",
    "竞赛",
    "比赛",
    "大赛",
    "一等奖",
    "二等奖",
    "三等奖",
    "荣誉",
    "项目",
    "完成",
    "成交",
    "优秀",
    "学生",
    "学员",
}

STOP_LINE_KEYWORDS = {
    "联系方式",
    "联系",
    "电话",
    "邮箱",
    "机构",
    "教育经历",
    "教育背景",
    "项目经历",
    "工作经历",
    "自我评价",
    "面试",
    "面销",
    "性别",
    "籍贯",
    "婚姻",
    "出生日期",
    "出生年月",
}

JOB_KEYWORDS = {
    "老师",
    "教师",
    "教练",
    "班主任",
    "顾问",
    "校长",
    "经理",
    "主管",
    "班主任",
    "数学",
    "物理",
    "英语",
    "语文",
    "化学",
    "政治",
    "历史",
    "地理",
    "生物",
    "体育",
    "课程",
    "招生",
}

NAME_SECTION_MARKERS = [
    "姓名",
    "姓 名",
    "基本信息",
    "个人信息",
    "个人简介",
    "联系信息",
    "contact",
    "contacts",
    "contact information",
    "personal information",
    "profile",
]
NAME_SECTION_MARKERS_LOWER = [marker.lower() for marker in NAME_SECTION_MARKERS]

LINE_FORBIDDEN_KEYWORDS = {
    "政治面貌",
    "婚姻状况",
    "兴趣爱好",
    "掌握技能",
    "技能特长",
    "资格证书",
    "证书",
    "荣誉奖项",
    "求职意向",
    "职业目标",
    "薪资",
    "期望",
    "现居地",
    "现住址",
    "联系地址",
    "家庭住址",
    "教育经历",
    "教育背景",
    "毕业院校",
    "毕业学校",
    "毕业时间",
    "工作经历",
    "项目经历",
    "项目经验",
    "自我评价",
    "自我介绍",
    "个人简介",
    "亮点",
    "摘要",
    "优势",
    "劣势",
    "兴趣特长",
    "特长",
    "证件",
    "身份证",
    "护照",
    "社保",
    "紧急联系人",
    "照片",
    "形象照",
    "作品集",
    "链接",
    "作品链接",
    "公众号",
    "网站",
    "平台",
    "掌握工具",
    "软件技能",
    "职业标签",
    "核心竞争力",
    "政治立场",
    "党员",
    "团员",
    "群众",
    "家庭成员",
    "个人情况",
    "人际关系",
    "特性描述",
    "性格",
    "MBTI",
    "血型",
    "健康状况",
}

COMMON_SURNAME_TOKENS = """
赵 钱 孙 李 周 吴 郑 王 冯 陈 褚 卫 蒋 沈 韩 杨 朱 秦 尤 许 何 吕 施 张
孔 曹 严 华 金 魏 陶 姜 戚 谢 邹 喻 柏 水 窦 章 云 苏 潘 葛 奚 范 彭 郎 鲁
韦 昌 马 苗 凤 花 方 俞 任 袁 柳 酆 鲍 史 唐 费 廉 岑 薛 雷 贺 倪 汤 滕
殷 罗 毕 郝 邬 安 常 乐 于 时 傅 皮 卞 齐 康 伍 余 元 卜 顾 孟 平 黄 和
穆 萧 尹 姚 邵 湛 汪 祁 毛 禹 狄 米 贝 明 臧 计 伏 成 戴 谈 宋 茅 庞 熊
纪 舒 屈 项 祝 董 梁 杜 阮 蓝 闵 席 季 麻 强 贾 路 娄 危 江 童 颜 郭 梅
盛 林 刁 钟 徐 邱 骆 高 夏 蔡 田 樊 胡 凌 霍 虞 万 支 柯 昝 管 卢 莫 经
房 裘 缪 干 解 应 宗 丁 宣 贲 邓 郁 单 杭 洪 包 诸 左 石 崔 吉 钮 龚 程
嵇 邢 滑 裴 陆 荣 翁 荀 羊 於 惠 甄 曲 家 封 芮 羿 储 靳 汲 邴 糜 松 井
段 富 巫 乌 焦 巴 弓 牧 隗 山 谷 车 侯 宓 蓬 全 郗 班 仰 秋 仲 伊 宫 宁 仇
栾 暴 甘 钭 厉 戎 祖 武 符 刘 景 詹 束 龙 叶 幸 司 韶 郜 黎 蓟 薄 印 宿
白 怀 蒲 邰 从 鄂 索 咸 籍 赖 卓 蔺 屠 蒙 池 乔 阴 鬱 胥 能 苍 双 闻 莘
党 翟 谭 贡 劳 逄 姬 申 扶 堵 冉 宰 郦 雍 郤 璩 桑 桂 濮 牛 寿 通 边 扈
燕 冀 郏 浦 尚 农 温 别 庄 晏 柴 瞿 阎 充 慕 连 茹 习 宦 艾 鱼 容 向 古
易 慎 戈 廖 庚 终 暨 居 衡 步 都 耿 满 弘 匡 国 文 寇 广 禄 阙 东 欧 殳
沃 利 蔚 越 夔 隆 师 巩 厍 聂 晁 勾 敖 融 冷 訾 辛 阚 那 简 饶 空 曾 毋
沙 乜 养 鞠 须 丰 巢 关 蒯 相 查 后 荆 红 游 竺 权 逯 盖 益 桓 公
万俟 司马 上官 欧阳 夏侯 诸葛 闻人 东方 赫连 皇甫 尉迟 公羊 澹台 公冶
宗政 濮阳 淳于 单于 太叔 申屠 公孙 仲孙 轩辕 令狐 钟离 宇文 长孙 慕容
鲜于 闾丘 司徒 司空 亓官 司寇 仉 督 子车 颛孙 端木 巫马 公西 漆雕 乐正
壤驷 公良 拓跋 夹谷 宰父 谷梁 段干 百里 东郭 南门 呼延 羊舌 微生 梁丘
左丘 东门 西门 南宫 第五 公仪 梁仲 公户 公玉 公仲 公上 公门 公山 公坚
谷利 谷利
"""
COMMON_SURNAME_CHARS = {token for token in COMMON_SURNAME_TOKENS.split() if len(token) == 1}
COMMON_DOUBLE_SURNAMES = {token for token in COMMON_SURNAME_TOKENS.split() if len(token) > 1}
COMMON_SURNAME_CHARS.update({"南", "付", "路", "那", "答", "雍", "覃"})

CONTACT_LINE_HINTS = {
    "手机",
    "电话",
    "mobile",
    "phone",
    "tel",
    "邮箱",
    "email",
    "mail",
    "wechat",
    "微信",
}

TITLE_SUFFIXES = (
    "老师",
    "同学",
    "先生",
    "小姐",
    "主管",
    "经理",
    "顾问",
    "总监",
)

IDENTITY_HINTS = {
    "姓名",
    "name",
    "Name",
    "性别",
    "男",
    "女",
    "年龄",
    "岁",
    "出生",
    "籍贯",
    "户籍",
    "民族",
    "住址",
    "地址",
    "联系方式",
    "个人信息",
    "基本信息",
}

CONTACT_PHONE_PATTERN = re.compile(r"(?:\+?86[-\s.]*)?(1[3-9]\d{9})")
CONTACT_EMAIL_PATTERN = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
CONTACT_TAIL_SPLIT = re.compile(
    r"(?:男|女|性别|电话|手机|邮箱|Email|mail|Mail|phone|Phone|联系方式|contact|Contact)",
    re.IGNORECASE,
)

FILENAME_NAME_PATTERNS = [
    re.compile(r"(?:姓名|name|Name)[\s:_-]*([\u4e00-\u9fa5·]{2,6})"),
    re.compile(r"([\u4e00-\u9fa5·]{2,4})\s*(?:\d{1,2}年)"),
    re.compile(r"([\u4e00-\u9fa5·]{2,4})(?:简历|求职|面试|作品)"),
]

CITY_WORDS = {"北京", "上海", "广州", "深圳", "杭州", "西安", "成都", "重庆", "苏州", "南京"}

ENGLISH_NAME_ALLOWED = re.compile(r"^[A-Za-z][A-Za-z\s\.\-]{1,30}$")
MAX_NAME_SCAN_LINES = 120
_LLM_CLIENT = None
_LLM_CFG = None


def _looks_like_suspicious_name(token: str) -> bool:
    if not token:
        return False
    if any(keyword in token for keyword in SUSPICIOUS_NAME_KEYWORDS):
        return True
    return bool(re.search(r"\d", token))


def _is_valid_name(token: str) -> bool:
    if not token:
        return False
    token = token.strip().strip(",.;:|")
    if not token:
        return False
    # 允许英文姓名
    if ENGLISH_NAME_ALLOWED.match(token):
        parts = [seg for seg in re.split(r"[\s\-]+", token) if seg]
        if 1 <= len(parts) <= 3 and all(part[0].isalpha() and part[0].isupper() for part in parts):
            return True
        return False
    # 中文姓名校验
    normalized = token.replace("·", "")
    if token in NAME_STOP_WORDS or token in CITY_WORDS:
        return False
    for substr in NAME_DISALLOWED_SUBSTRINGS:
        if substr in token:
            return False
    if _looks_like_suspicious_name(token):
        return False
    if not normalized:
        return False
    if not re.fullmatch(r"[\u4e00-\u9fa5]+", normalized):
        return False
    length = len(normalized)
    if length < 2 or length > 4:
        return False
    prefix2 = normalized[:2]
    has_double_surname = prefix2 in COMMON_DOUBLE_SURNAMES
    if length == 3 and has_double_surname:
        return True
    if length >= 4 and has_double_surname:
        return True
    first_char = normalized[0]
    if first_char not in COMMON_SURNAME_CHARS:
        return False
    return True


def _normalize_name_token(token: str) -> str:
    if not token:
        return ""
    token = token.strip().strip("，,。.;:|/\\-")
    token = re.sub(r"\s+", " ", token)
    return token


def _clean_candidate_token(token: str) -> str:
    if not token:
        return ""
    token = CONTACT_TAIL_SPLIT.split(token)[0]
    token = re.split(r"[，,。；;、/|]", token)[0]
    for suffix in TITLE_SUFFIXES:
        if token.endswith(suffix) and len(token) > len(suffix):
            stripped = token[: -len(suffix)]
            if len(stripped) >= 2:
                token = stripped
            break
    return _normalize_name_token(token)


def _line_has_identity_signal(line: str) -> bool:
    if not line:
        return False
    lowered = line.lower()
    if re.search(r"(姓\s*名|姓名|name)", lowered):
        return True
    if re.search(r"(男|女)", line):
        return True
    if re.search(r"\d{2}岁|\d{4}年", line):
        return True
    return any(hint.lower() in lowered for hint in IDENTITY_HINTS)


def _line_is_mostly_candidate(line: str, candidate: str) -> bool:
    if not line or not candidate:
        return False
    compact = re.sub(r"[\s，,。：:；;、/\\|_()-]+", "", line)
    return compact == candidate


def _looks_like_filename_name(token: str, allow_titles: bool = False) -> bool:
    if not token:
        return False
    if token in NAME_STOP_WORDS or token in CITY_WORDS:
        return False
    disallowed = NAME_DISALLOWED_SUBSTRINGS
    if allow_titles:
        disallowed = {item for item in NAME_DISALLOWED_SUBSTRINGS if item not in TITLE_SUFFIXES}
    if any(substr in token for substr in disallowed):
        return False
    return bool(re.fullmatch(r"[\u4e00-\u9fa5·]{2,4}", token))


def _prepare_lines(text: str) -> List[str]:
    if not text:
        return []
    lines: List[str] = []
    for raw in text.splitlines():
        clean = raw.strip()
        if not clean:
            continue
        lines.append(clean)
        if len(lines) >= MAX_NAME_SCAN_LINES:
            break
    return lines


def _contains_contact_hint(line: str) -> bool:
    if not line:
        return False
    lowered = line.lower()
    if CONTACT_PHONE_PATTERN.search(line) or CONTACT_EMAIL_PATTERN.search(line):
        return True
    return any(hint in line or hint in lowered for hint in CONTACT_LINE_HINTS)


def _extract_candidate_from_line(line: str) -> str:
    if not line:
        return ""
    normalized_line = line.strip().lower()
    if normalized_line in NAME_SECTION_MARKERS_LOWER:
        return ""
    if any(keyword in line for keyword in LINE_FORBIDDEN_KEYWORDS):
        if not re.search(r"(姓\s*名|姓名|name|Name)", line):
            return ""
    if (":" in line or "：" in line) and any(keyword in line for keyword in SUSPICIOUS_NAME_KEYWORDS):
        return ""
    identity_signal = _line_has_identity_signal(line)
    labeled_match = re.search(r"(?:姓\s*名|姓名|name|Name)[:：\s]*([A-Za-z\u4e00-\u9fa5·\s]{2,30})", line)
    if labeled_match:
        candidate = _clean_candidate_token(labeled_match.group(1))
        if _is_valid_name(candidate):
            return candidate
    for size in range(6, 1, -1):
        for match in re.finditer(r"[\u4e00-\u9fa5·]{" + str(size) + r"}", line):
            candidate = _clean_candidate_token(match.group(0))
            if _is_valid_name(candidate) and (identity_signal or _line_is_mostly_candidate(line, candidate)):
                return candidate
    # 英文姓名在整行内
    english_match = ENGLISH_NAME_ALLOWED.search(line)
    if english_match:
        candidate = _normalize_name_token(english_match.group(0))
        if candidate.lower() in NAME_SECTION_MARKERS_LOWER:
            return ""
        if _is_valid_name(candidate) and (identity_signal or _line_is_mostly_candidate(line, candidate)):
            return candidate
    return ""


def _extract_from_marked_sections(lines: List[str]) -> str:
    for idx, line in enumerate(lines):
        compact = line.replace(" ", "").lower()
        if any(marker in compact for marker in NAME_SECTION_MARKERS):
            candidate = _extract_candidate_from_line(line)
            if candidate:
                return candidate
            for look_ahead in range(1, 3):
                if idx + look_ahead < len(lines):
                    candidate = _extract_candidate_from_line(lines[idx + look_ahead])
                    if candidate:
                        return candidate
    return ""


def _extract_from_contact_blocks(lines: List[str]) -> str:
    for idx, line in enumerate(lines):
        if not _contains_contact_hint(line):
            continue
        window = []
        for offset in (-2, -1, 0, 1, 2):
            pos = idx + offset
            if 0 <= pos < len(lines):
                window.append((pos, lines[pos]))
        for pos, candidate_line in window:
            candidate = _extract_candidate_from_line(candidate_line)
            if candidate and (pos <= idx or _line_is_mostly_candidate(candidate_line, candidate)):
                return candidate
        for pos, candidate_line in window:
            tokens = re.findall(r"[\u4e00-\u9fa5·]{2,4}", candidate_line)
            identity_signal = _line_has_identity_signal(candidate_line)
            for token in tokens:
                cleaned = _clean_candidate_token(token)
                if not _is_valid_name(cleaned):
                    continue
                allow_line = identity_signal or _line_is_mostly_candidate(candidate_line, cleaned)
                if allow_line:
                    return cleaned
    return ""


def _get_llm_client():
    global _LLM_CLIENT, _LLM_CFG
    if _LLM_CLIENT is not None or get_client_and_cfg is None:
        return _LLM_CLIENT, _LLM_CFG
    try:
        _LLM_CLIENT, _LLM_CFG = get_client_and_cfg()
    except Exception:
        _LLM_CLIENT = None
        _LLM_CFG = None
    return _LLM_CLIENT, _LLM_CFG


def _llm_extract_name_from_text(text: str) -> str:
    if not text or not text.strip():
        return ""
    if chat_completion is None:
        return ""
    client, cfg = _get_llm_client()
    if not client or not cfg or not getattr(cfg, "api_key", None):
        return ""
    snippet = textwrap.shorten(text.strip(), width=1800, placeholder=" ...")
    prompt = textwrap.dedent(
        f"""
        你是一名简历解析助手。请从以下正文片段中提取候选人的真实姓名。
        必须遵守：
        1. 只能使用正文片段中的信息，不得根据文件名或常识猜测。
        2. 不得编造姓名，如果未找到，请返回 NONE。
        3. 如果姓名出现多次，选择最明显的真实姓名。

        正文片段：
        <<<{snippet}>>>

        请严格按照格式输出：NAME: <姓名或NONE>
        """
    ).strip()
    try:
        response = chat_completion(
            client,
            cfg,
            messages=[
                {"role": "system", "content": "You extract real candidate names from resume text and never guess."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.0,
            max_tokens=32,
            stop=["\n\n"],
        )
        text_resp = response["choices"][0]["message"]["content"].strip()
        match = re.search(r"NAME[:：]\s*([A-Za-z\u4e00-\u9fa5·\s]{2,40}|NONE)", text_resp, re.IGNORECASE)
        if not match:
            return ""
        extracted = match.group(1).strip()
        if extracted.upper() == "NONE":
            return ""
        extracted = _normalize_name_token(extracted)
        if _is_valid_name(extracted) and extracted in text:
            return extracted
    except Exception:
        return ""
    return ""

def _extract_name_from_text(text: str) -> str:
    if not text:
        return ""
    for pattern in NAME_PATTERNS:
        match = pattern.search(text)
        if match:
            candidate = _clean_candidate_token(match.group(1))
            if _is_valid_name(candidate):
                return candidate
    lines = _prepare_lines(text)
    candidate = _extract_from_marked_sections(lines)
    if candidate:
        return candidate
    candidate = _extract_from_contact_blocks(lines)
    if candidate:
        return candidate
    candidates: List[Tuple[str, float]] = []

    def _line_score(line: str) -> float:
        score = 0.0
        if "姓名" in line or "Name" in line:
            score += 3
        if re.search(r"(男|女)", line):
            score += 1
        if any(keyword in line for keyword in ("基本信息", "个人信息", "联系方式", "联系方式：")):
            score += 1.5
        if re.search(r"\d{2}岁|\d{4}年", line):
            score += 0.5
        return score

    for line in lines[:40]:
        if any(keyword in line for keyword in STOP_LINE_KEYWORDS):
            continue
        tokens = re.findall(r"[\u4e00-\u9fa5·]{2,6}", line)
        score = _line_score(line)
        identity_signal = _line_has_identity_signal(line)
        for token in tokens:
            cleaned = _clean_candidate_token(token)
            if _is_valid_name(cleaned) and (identity_signal or _line_is_mostly_candidate(line, cleaned)):
                candidates.append((cleaned, score))

    if candidates:
        candidates.sort(key=lambda x: (x[1], len(x[0])), reverse=True)
        best = candidates[0]
        if best[1] >= 1:
            return best[0]
    return ""


def _extract_name_from_filename(filename: str, allow_titles: bool = False) -> str:
    raw_stem = Path(filename).stem
    # 常见形式：...】姓名_x年 / ...]姓名
    direct_match = re.search(r"[】\]]\s*([\u4e00-\u9fa5·]{2,6})", raw_stem)
    if direct_match:
        candidate = direct_match.group(1)
        if _is_valid_name(candidate) or _looks_like_filename_name(candidate, allow_titles=allow_titles):
            return candidate

    stem = raw_stem
    # 去除首尾括号包裹的标签
    stem = re.sub(r"^[\[\(（【].*?[\]\)）】]", "", stem).strip()
    stem = re.sub(r"[\(\（][^)\）]*[\)\）]", " ", stem)
    stem = stem.replace("_", " ").replace("-", " ").replace("（", " ").replace("）", " ")
    stem = re.sub(r"[【】\[\]\(\)（）]", " ", stem)
    stem = re.sub(r"\s+", " ", stem).strip()
    stem = stem.lstrip("】] ")
    tail_match = re.search(r"([\u4e00-\u9fa5·]{2,4})\s*(?:\d{2,4})?(?:版)?$", stem)
    if tail_match:
        candidate = tail_match.group(1)
        if _is_valid_name(candidate) or _looks_like_filename_name(candidate, allow_titles=allow_titles):
            return candidate
    for pattern in FILENAME_NAME_PATTERNS:
        match = pattern.search(stem)
        if match:
            candidate = _normalize_name_token(match.group(1))
            if _is_valid_name(candidate) or _looks_like_filename_name(candidate, allow_titles=allow_titles):
                return candidate

    # 评分策略：越靠近末尾、后接“年/岁/简历”等词得分越高
    tokens = list(re.finditer(r"[\u4e00-\u9fa5·]{2,6}", stem))
    scored: List[Tuple[str, float]] = []
    for match in tokens:
        token = match.group(0)
        if not _is_valid_name(token):
            continue
        start = match.start()
        end = match.end()
        score = 1.0
        if len(stem) - end < 8:
            score += 1.5
        if re.match(r"(?:\d{1,2}年|\d{1,2}岁)", stem[end:end + 3]):
            score += 1.5
        if any(keyword in stem[max(0, start - 4):end + 4] for keyword in ("姓名", "name", "Name")):
            score += 2
        prev_chunk = stem[max(0, start - 6):start]
        if any(keyword in prev_chunk for keyword in JOB_KEYWORDS):
            score -= 1.5
        scored.append((token, score))
    if scored:
        scored.sort(key=lambda x: x[1], reverse=True)
        top_candidate, top_score = scored[0]
        if top_score >= 2 and (_is_valid_name(top_candidate) or _looks_like_filename_name(top_candidate, allow_titles=allow_titles)):
            return top_candidate
    english_match = re.search(r"([A-Za-z][A-Za-z\s\.\-]{1,30})\s*(?:\d{1,3})?$", stem)
    if english_match:
        candidate = _normalize_name_token(english_match.group(1))
        if _is_valid_name(candidate):
            return candidate
    return ""


def infer_candidate_name(text: str, filename: str) -> str:
    name = _extract_name_from_text(text)
    if name and not _looks_like_suspicious_name(name):
        return name
    llm_name = _llm_extract_name_from_text(text)
    if llm_name and not _looks_like_suspicious_name(llm_name):
        return llm_name
    fallback = _extract_name_from_filename(filename)
    if fallback and not _looks_like_suspicious_name(fallback):
        return fallback
    relaxed = _extract_name_from_filename(filename, allow_titles=True)
    if relaxed:
        return relaxed
    return ""


def extract_pdf_text(path: Path) -> str:
    """提取PDF文本,使用多种方法确保成功"""
    text = ""
    
    # 方法1: 使用 PyMuPDF (fitz) - 最常用且快速
    try:
        doc = fitz.open(str(path))
        raw_parts = []
        for page_num in range(len(doc)):
            page = doc[page_num]
            # 优先使用标准文本提取
            page_text = page.get_text("text") or ""
            # 如果标准方法失败,尝试使用blocks方法
            if not page_text.strip():
                blocks = page.get_text("blocks")
                if blocks:
                    block_texts = [
                        str(block[4])
                        for block in blocks
                        if len(block) > 4 and isinstance(block[4], str)
                    ]
                    page_text = "\n".join(block_texts)
            # 再尝试原始字典
            if not page_text.strip():
                raw_dict = page.get_text("rawdict")
                if isinstance(raw_dict, dict):
                    page_text = json.dumps(raw_dict, ensure_ascii=False)
                elif raw_dict:
                    page_text = str(raw_dict)
            raw_parts.append(page_text)
        doc.close()
        text = _clean_text("\n".join(raw_parts))
        if len(text.strip()) >= 50:
            return text
    except Exception:
        pass
    
    # 方法2: 使用 pdfplumber (如果可用) - 对某些PDF格式更有效
    if pdfplumber:
        try:
            with pdfplumber.open(str(path)) as pdf:
                pages = []
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    if not page_text.strip():
                        tables = page.extract_tables()
                        if tables:
                            for table in tables:
                                for row in table:
                                    if row:
                                        row_text = " ".join([str(cell) if cell else "" for cell in row])
                                        page_text += row_text + "\n"
                    pages.append(page_text)
            new_text = _clean_text("\n".join(pages))
            if len(new_text.strip()) >= len(text.strip()):
                text = new_text
            if len(text.strip()) >= 50:
                return text
        except Exception:
            pass
    
    return text


def ocr_pdf(path: Path) -> str:
    """使用OCR提取PDF文本(用于扫描版PDF)"""
    try:
        poppler_path = os.getenv("POPPLER_PATH")
        # 尝试转换PDF为图片
        pages = convert_from_path(str(path), dpi=300, poppler_path=poppler_path)
        text_chunks: List[str] = []
        for page in pages:
            try:
                # 尝试中文OCR,如果失败则使用英文
                page_text = pytesseract.image_to_string(page, lang="chi_sim+eng")
                if not page_text.strip():
                    page_text = pytesseract.image_to_string(page, lang="eng")
                text_chunks.append(page_text)
            except Exception:
                # 如果OCR失败,尝试只使用英文
                try:
                    page_text = pytesseract.image_to_string(page, lang="eng")
                    text_chunks.append(page_text)
                except Exception:
                    text_chunks.append("")
        result = _clean_text("\n".join(text_chunks))
        return result
    except Exception as e:
        # OCR失败,返回空字符串
        return ""


def parse_pdf(path: Path) -> Tuple[str, Dict[str, str]]:
    """解析PDF文件,优先使用文本提取,失败时使用OCR"""
    # 首先尝试直接提取文本
    text = extract_pdf_text(path)
    
    # 如果提取的文本太少(可能是扫描版PDF),尝试OCR
    if len(text.strip()) < 100:
        ocr_text = ocr_pdf(path)
        if ocr_text and len(ocr_text.strip()) > len(text.strip()):
            text = ocr_text
    
    # 如果仍然没有文本,尝试使用PyMuPDF的其他方法
    if not text.strip():
        try:
            doc = fitz.open(str(path))
            raw_parts = []
            for page_num in range(len(doc)):
                page = doc[page_num]
                # 尝试不同的文本提取方法
                page_text = page.get_text("text") or ""
                if not page_text.strip():
                    # 尝试获取原始文本
                    page_text = page.get_text("rawdict") or ""
                    if isinstance(page_text, dict):
                        page_text = str(page_text)
                raw_parts.append(page_text)
            doc.close()
            text = "\n".join(raw_parts)
            text = _clean_text(text)
        except Exception:
            pass
    
    # 提取联系信息
    contacts = extract_contacts(text)
    
    return text, contacts


def parse_docx(path: Path) -> Tuple[str, Dict[str, str]]:
    text = ""
    if docx2txt:
        try:
            text = docx2txt.process(str(path)) or ""
        except Exception:
            text = ""
    if (not text or len(text.strip()) < 20) and Document:
        try:
            doc = Document(str(path))
            parts = []
            for para in doc.paragraphs:
                if para.text:
                    parts.append(para.text)
            for table in getattr(doc, "tables", []):
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text]
                    if cells:
                        parts.append(" ".join(cells))
            text = "\n".join(parts)
        except Exception:
            pass
    if not text or len(text.strip()) < 20:
        text = _extract_docx_via_xml(path)
    cleaned = _clean_text(text)
    return cleaned, extract_contacts(cleaned)


def _extract_docx_via_xml(path: Path) -> str:
    try:
        with zipfile.ZipFile(path) as zf:
            xml_bytes = zf.read("word/document.xml")
        root = ElementTree.fromstring(xml_bytes)
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        paragraphs = []
        for para in root.findall(".//w:p", ns):
            texts = [node.text for node in para.findall(".//w:t", ns) if node.text]
            if texts:
                paragraphs.append("".join(texts))
        return "\n".join(paragraphs)
    except Exception:
        return ""


def parse_txt(path: Path) -> Tuple[str, Dict[str, str]]:
    raw = path.read_bytes()
    text = _detect_encoding_and_read(raw)
    return _clean_text(text), extract_contacts(text)


def parse_image(path: Path) -> Tuple[str, Dict[str, str]]:
    text = ""
    try:
        with Image.open(str(path)) as img:
            text = pytesseract.image_to_string(img, lang="chi_sim+eng")
    except Exception:
        text = ""
    return _clean_text(text), extract_contacts(text)


def parse_one_to_text(path: Path) -> Tuple[str, Dict[str, str]]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(path)
    if suffix == ".docx":
        return parse_docx(path)
    if suffix == ".txt":
        return parse_txt(path)
    if suffix in {".jpg", ".jpeg", ".png"}:
        return parse_image(path)
    return "", {"email": "", "phone": ""}


def save_uploaded_to_tmp(uploaded_file, out_dir: Path) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    filename = Path(uploaded_file.name).name
    tmp_path = out_dir / filename
    tmp_path.write_bytes(uploaded_file.getbuffer())
    return tmp_path


def parse_uploaded_files_to_df(files: List, max_chars: int = 20000) -> pd.DataFrame:
    rows = []
    out_dir = Path("data/uploads")
    out_dir.mkdir(parents=True, exist_ok=True)

    cid = 1
    for uploaded in files:
        suffix = Path(uploaded.name).suffix.lower()
        if suffix not in SUPPORTED_EXT:
            continue
        tmp_path = save_uploaded_to_tmp(uploaded, out_dir)
        
        # 解析文件
        text, contacts = parse_one_to_text(tmp_path)
        
        # 如果解析失败(文本为空或太短),尝试其他方法
        if not text.strip() or len(text.strip()) < 50:
            # 对于PDF文件,尝试更激进的解析方法
            if suffix == ".pdf":
                try:
                    # 再次尝试使用PyMuPDF,使用不同的参数
                    doc = fitz.open(str(tmp_path))
                    raw_parts = []
                    for page_num in range(len(doc)):
                        page = doc[page_num]
                        # 尝试多种文本提取方法
                        page_text = page.get_text("text") or ""
                        if not page_text.strip():
                            # 尝试使用blocks方法
                            blocks = page.get_text("blocks")
                            if blocks:
                                page_text = "\n".join([block[4] for block in blocks if len(block) > 4])
                        raw_parts.append(page_text)
                    doc.close()
                    text = "\n".join(raw_parts)
                    text = _clean_text(text)
                except Exception:
                    pass
        
        # 限制文本长度
        text = text[:max_chars] if text else ""
        candidate_name = infer_candidate_name(text, tmp_path.name)
        
        rows.append(
            {
                "candidate_id": cid,
                "file": tmp_path.name,
                "name": candidate_name,
                "resume_text": text,
                "text_len": len(text),
                "email": contacts.get("email", ""),
                "phone": contacts.get("phone", ""),
            }
        )
        cid += 1

    df = pd.DataFrame(rows)
    if df.empty:
        return pd.DataFrame(columns=["candidate_id", "file", "name", "resume_text", "text_len", "email", "phone"])

    source_columns = ["resume_text", "text", "full_text", "content", "parsed_text"]

    def _coalesce_resume_text(row):
        for col in source_columns:
            if col in row and isinstance(row[col], str) and row[col].strip():
                return row[col].strip()
        for col in source_columns:
            if col in row and row[col]:
                return str(row[col]).strip()
        return ""

    df["resume_text"] = df.apply(_coalesce_resume_text, axis=1)
    df["resume_text"] = df["resume_text"].fillna("").astype(str)
    df["text_len"] = df["resume_text"].str.len()

    drop_columns = [c for c in ["text", "full_text", "content", "parsed_text"] if c in df.columns]
    if drop_columns:
        df = df.drop(columns=drop_columns)

def ocr_pdf(path: Path) -> str:
    """使用OCR提取PDF文本(用于扫描版PDF)"""
    try:
        poppler_path = os.getenv("POPPLER_PATH")
        # 尝试转换PDF为图片
        pages = convert_from_path(str(path), dpi=300, poppler_path=poppler_path)
        text_chunks: List[str] = []
        for page in pages:
            try:
                # 尝试中文OCR,如果失败则使用英文
                page_text = pytesseract.image_to_string(page, lang="chi_sim+eng")
                if not page_text.strip():
                    page_text = pytesseract.image_to_string(page, lang="eng")
                text_chunks.append(page_text)
            except Exception:
                # 如果OCR失败,尝试只使用英文
                try:
                    page_text = pytesseract.image_to_string(page, lang="eng")
                    text_chunks.append(page_text)
                except Exception:
                    text_chunks.append("")
        result = _clean_text("\n".join(text_chunks))
        return result
    except Exception as e:
        # OCR失败,返回空字符串
        return ""


def parse_pdf(path: Path) -> Tuple[str, Dict[str, str]]:
    """解析PDF文件,优先使用文本提取,失败时使用OCR"""
    # 首先尝试直接提取文本
    text = extract_pdf_text(path)
    
    # 如果提取的文本太少(可能是扫描版PDF),尝试OCR
    if len(text.strip()) < 100:
        ocr_text = ocr_pdf(path)
        if ocr_text and len(ocr_text.strip()) > len(text.strip()):
            text = ocr_text
    
    # 如果仍然没有文本,尝试使用PyMuPDF的其他方法
    if not text.strip():
        try:
            doc = fitz.open(str(path))
            raw_parts = []
            for page_num in range(len(doc)):
                page = doc[page_num]
                # 尝试不同的文本提取方法
                page_text = page.get_text("text") or ""
                if not page_text.strip():
                    # 尝试获取原始文本
                    page_text = page.get_text("rawdict") or ""
                    if isinstance(page_text, dict):
                        page_text = str(page_text)
                raw_parts.append(page_text)
            doc.close()
            text = "\n".join(raw_parts)
            text = _clean_text(text)
        except Exception:
            pass
    
    # 提取联系信息
    contacts = extract_contacts(text)
    
    return text, contacts


def parse_docx(path: Path) -> Tuple[str, Dict[str, str]]:
    text = ""
    if docx2txt:
        try:
            text = docx2txt.process(str(path)) or ""
        except Exception:
            text = ""
    if (not text or len(text.strip()) < 20) and Document:
        try:
            doc = Document(str(path))
            parts = []
            for para in doc.paragraphs:
                if para.text:
                    parts.append(para.text)
            for table in getattr(doc, "tables", []):
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text]
                    if cells:
                        parts.append(" ".join(cells))
            text = "\n".join(parts)
        except Exception:
            pass
    if not text or len(text.strip()) < 20:
        text = _extract_docx_via_xml(path)
    cleaned = _clean_text(text)
    return cleaned, extract_contacts(cleaned)


def _extract_docx_via_xml(path: Path) -> str:
    try:
        with zipfile.ZipFile(path) as zf:
            xml_bytes = zf.read("word/document.xml")
        root = ElementTree.fromstring(xml_bytes)
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        paragraphs = []
        for para in root.findall(".//w:p", ns):
            texts = [node.text for node in para.findall(".//w:t", ns) if node.text]
            if texts:
                paragraphs.append("".join(texts))
        return "\n".join(paragraphs)
    except Exception:
        return ""


def parse_txt(path: Path) -> Tuple[str, Dict[str, str]]:
    raw = path.read_bytes()
    text = _detect_encoding_and_read(raw)
    return _clean_text(text), extract_contacts(text)


def parse_image(path: Path) -> Tuple[str, Dict[str, str]]:
    text = ""
    try:
        with Image.open(str(path)) as img:
            text = pytesseract.image_to_string(img, lang="chi_sim+eng")
    except Exception:
        text = ""
    return _clean_text(text), extract_contacts(text)


def parse_one_to_text(path: Path) -> Tuple[str, Dict[str, str]]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(path)
    if suffix == ".docx":
        return parse_docx(path)
    if suffix == ".txt":
        return parse_txt(path)
    if suffix in {".jpg", ".jpeg", ".png"}:
        return parse_image(path)
    return "", {"email": "", "phone": ""}


def save_uploaded_to_tmp(uploaded_file, out_dir: Path) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    filename = Path(uploaded_file.name).name
    tmp_path = out_dir / filename
    tmp_path.write_bytes(uploaded_file.getbuffer())
    return tmp_path


def parse_uploaded_files_to_df(files: List, max_chars: int = 20000) -> pd.DataFrame:
    rows = []
    out_dir = Path("data/uploads")
    out_dir.mkdir(parents=True, exist_ok=True)

    cid = 1
    for uploaded in files:
        suffix = Path(uploaded.name).suffix.lower()
        if suffix not in SUPPORTED_EXT:
            continue
        tmp_path = save_uploaded_to_tmp(uploaded, out_dir)
        
        # 解析文件
        text, contacts = parse_one_to_text(tmp_path)
        
        # 如果解析失败(文本为空或太短),尝试其他方法
        if not text.strip() or len(text.strip()) < 50:
            # 对于PDF文件,尝试更激进的解析方法
            if suffix == ".pdf":
                try:
                    # 再次尝试使用PyMuPDF,使用不同的参数
                    doc = fitz.open(str(tmp_path))
                    raw_parts = []
                    for page_num in range(len(doc)):
                        page = doc[page_num]
                        # 尝试多种文本提取方法
                        page_text = page.get_text("text") or ""
                        if not page_text.strip():
                            # 尝试使用blocks方法
                            blocks = page.get_text("blocks")
                            if blocks:
                                page_text = "\n".join([block[4] for block in blocks if len(block) > 4])
                        raw_parts.append(page_text)
                    doc.close()
                    text = "\n".join(raw_parts)
                    text = _clean_text(text)
                except Exception:
                    pass
        
        # 限制文本长度
        text = text[:max_chars] if text else ""
        candidate_name = infer_candidate_name(text, tmp_path.name)
        
        rows.append(
            {
                "candidate_id": cid,
                "file": tmp_path.name,
                "name": candidate_name,
                "resume_text": text,
                "text_len": len(text),
                "email": contacts.get("email", ""),
                "phone": contacts.get("phone", ""),
            }
        )
        cid += 1

    df = pd.DataFrame(rows)
    if df.empty:
        return pd.DataFrame(columns=["candidate_id", "file", "name", "resume_text", "text_len", "email", "phone"])

    source_columns = ["resume_text", "text", "full_text", "content", "parsed_text"]

    def _coalesce_resume_text(row):
        for col in source_columns:
            if col in row and isinstance(row[col], str) and row[col].strip():
                return row[col].strip()
        for col in source_columns:
            if col in row and row[col]:
                return str(row[col]).strip()
        return ""

    df["resume_text"] = df.apply(_coalesce_resume_text, axis=1)
    df["resume_text"] = df["resume_text"].fillna("").astype(str)
    df["text_len"] = df["resume_text"].str.len()

    drop_columns = [c for c in ["text", "full_text", "content", "parsed_text"] if c in df.columns]
    if drop_columns:
        df = df.drop(columns=drop_columns)

    return df

